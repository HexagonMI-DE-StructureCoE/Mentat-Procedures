from py_mentat import * 

import os
import sys 

import subprocess
current_path = os.path.dirname(os.path.abspath(__file__))
new_path = os.path.join(current_path, "..", "PYTHON_LIB")

print(new_path)

sys.path.append(new_path)


try:
    from docx import Document
    print("ok package is installed.")
except:

    python_path = sys.executable
    python_0 = os.path.dirname(python_path)
    if os.name == "nt":
        python_home = os.path.join(python_0,"..","..","python","WIN8664")
    else:
        # TODO:check
        python_home = os.path.join(python_0,"..","..","python","lx84")

    if os.name == "nt":
        fileo = open("installa.bat","w")
        fileo.write("@echo off\n")
        fileo.write("\n")
        fileo.write("set PYTHONPATH=..\PYTHON_LIB\n")
        fileo.write("set PYTHONHOME=%s\n" % python_home )
        fileo.write("\n")
        fileo.write("set PATH=%PYTHONHOME%\PyQt5;%PYTHONHOME%\;%PYTHONHOME%\DLLs;%PYTHONHOME%\Scripts\n")
        fileo.write("set PATH=%PATH%;c:\Windows;c:\Windows\System32;c:\Windows\Wbem\n")
        fileo.write("\n")
        fileo.write("\n")
        fileo.write("SET SUBDIR=%~dp0\n")
        fileo.write("cd /D %SUBDIR% \n")

        fileo.write("echo Do you want to execute the next command? (Y/N)\n")
        fileo.write("choice /c YN\n")
        fileo.write("\n")
        fileo.write("if errorlevel 2 (\n")
        fileo.write("    echo Command execution canceled.\n")
        fileo.write("exit")
        fileo.write(") else (\n")
        fileo.write("    echo Proceeding with the next command.\n")


        fileo.write("cd /D %s\n" % new_path)
        fileo.write("%PYTHONHOME%\\python -m pip install python-docx --target=..\PYTHON_LIB --no-user \n" )

        fileo.write("echo ....\n")
        fileo.write("echo Exit from Mentat and rerun the tool \n")


        fileo.write("pause")
        fileo.write("exit")
        fileo.write(")\n")

        fileo.close()



        cur_file_dir= os.path.dirname(os.path.abspath(__file__))
        install_location = os.path.join(cur_file_dir,"..","PYTHON_LIB")

        if not os.path.exists(install_location):
            print("creating folder")
            os.makedirs(install_location)
        
        if os.name == "nt":
            subprocess.Popen(['start', 'cmd', '/k', 'installa.bat'], shell=True)

    else:
        print("The python-docx package was not installed.")




from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def main():

    
    
    document = Document()

    
    modelname0 = py_get_string("filename()")
    modelname1 = modelname0.replace(".t16","").replace(".t19","")
    
    titolo = "Report Model\n" + modelname1
    titolo = titolo
    d = document.add_heading(titolo , 0)
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = document.add_paragraph('')
    p.add_run("\n")
    modelname = py_get_string('model_name()')
    p.add_run(f" Model name: {modelname} \n")

    py_send("*fill_view")
    py_send("*post_rewind")
    py_send("*post_off")
    py_send('*image_save_current "results.png" yes')
    
    if os.path.exists("model.png"):
        document.add_picture('model.png' , width=Inches(6.0) )
    
    #Numero elementi
    document.add_heading("Model description\n" , 1)
    p1 = document.add_paragraph('')
    nodi = py_get_int("nnodes()")
    str1 = " Number of Nodes       %g \n" % (nodi)
    p1.add_run(str1)
    elementi = py_get_int("nelements()")
    str1 = " Number of Elements %g \n" % (elementi)
    p1.add_run(str1)
    
      
    document.add_page_break()
    
    


    document.add_heading("Results description\n" , 1)
    p1 = document.add_paragraph('')
    #p1.add_run(stri)
    
    
    lista_res = ["Displacement"]    # , "Total Equivalent Plastic Strain","Equivalent Von Mises Stress"]
    
    py_send("*post_skip_to_last")
    py_send("*system_grid_display_off")
    
    py_send("*post_contour_bands")

    print(lista_res)
    
    for val1 in lista_res:
        py_send("*post_value %s" % val1 )
        py_send('*image_save_current "%s.png" yes' % val1)
        if os.path.exists("%s.png" % val1 ):
            p1 = document.add_paragraph('')
            p1.add_run("Displacement\n")
            document.add_picture('Displacement.png' , width=Inches(6.0) )
            document.add_page_break()
    
    filename = "report_" + str(modelname1) + "_results.docx"

    document.save(filename)
    
    try:
        #py_send("*system_command %s" % filename)
        subprocess.Popen(['start', 'winword', filename], shell=True)
        pass
    except:
        print("no opending file")

    py_prompt("Done")