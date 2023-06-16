from py_mentat import *

import os
import sys 

import subprocess
    
import install_module_docx

install_module_docx.installa_modulo()


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def main():
    document = Document()

    modelname0 = py_get_string("filename()")
    modelname1 = modelname0.replace(".mud","").replace(".mfd","")
    titolo = "Report Model\n" + modelname1
    titolo = titolo
    d = document.add_heading(titolo , 0)
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = document.add_paragraph('')
    p.add_run("\n")
    modelname = py_get_string('model_name()')
    p.add_run(f" Model name: {modelname} \n")

    py_send("*fill_view")
    py_send('*image_save_current "model.png" yes')
    
    if os.path.exists("model.png"):
        document.add_picture('model.png' , width=Inches(6.0) )
    
    #Numero elementi
    document.add_heading("Model description\n" , 1)
    p = document.add_paragraph('')
    nodi = py_get_int("nnodes()")
    str1 = " Number of Nodes       %g \n" % (nodi)
    p.add_run(str1)
    elementi = py_get_int("nelements()")
    str1 = " Number of Elements %g \n" % (elementi)
    p.add_run(str1)
      
    document.add_page_break()

    document.add_heading("Materials\n" , 1)
    p = document.add_paragraph('')
    m = py_get_int("nmaters()")
    p.add_run(f"Number of Materials {m} \n\n") 
    
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    #table.style = 'TableGrid'
    #table.style = 'ColorfulShading'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Material Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Young'
    hdr_cells[3].text = 'Poisson'
    hdr_cells[4].text = 'Yield'
    
    
    
    for i in range(1,m+1):
        
        sn = py_get_string("mater_name_index(%d)" % i)
        #p.add_run(f"  Material name   : {sn} \n")
        st = py_get_string("mater_type(%s)" % sn)
        #p.add_run(f"     Type              : {st} \n" )
        e = py_get_float("mater_par(%s,structural:youngs_modulus)" % sn)
        #p.add_run(f"     Youngs Modulus    : {e} \n")
        p1 = py_get_float("mater_par(%s,structural:poissons_ratio)" % sn)
        #p.add_run(f"      Poissons Ratio    : {p1} \n") 
        ys = py_get_float("mater_par(%s,structural:yield_stress)" % sn)
        #p.add_run(f"     Yield Stress      : {ys} \n")
        
        row_cells = table.add_row().cells
        #row_cells[0].style = "borderColor:red;background-color:gray"
        row_cells[0].text = str(sn)
        row_cells[1].text = str(st)
        row_cells[2].text = str(e)
        #p=row_cells[2].add_paragraph( str(e) )
        #p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        #row_cells[2].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[3].text = str(p1)
        #table.cell(3,i).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[4].text = str(ys)
        #table.cell(4,i).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    
    document.add_page_break()

    document.add_heading("Boundary Conditions\n" , 1)
    p = document.add_paragraph('')
    m = py_get_int("napplys()")
    p.add_run(f"Number of Boundary Conditions {m} \n\n")
    
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    #table.style = 'TableGrid'
    #table.style = 'ColorfulShading'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Boundary Condition'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Value'
    
    for i in range(1,m+1):
        sn = py_get_string("apply_name_index(%d)" % i)
        st = py_get_string("apply_type(%s)" % sn)
        so = py_get_string("apply_opt(%s,dof_values)" % sn)
        #str1 = "  Boundary Cond %14s  Type %19s   Values by: %s\n" % (sn,st,so)
        #p.add_run(str1)
        
        row_cells = table.add_row().cells
        #row_cells[0].style = "borderColor:red;background-color:gray"
        row_cells[0].text = str(sn)
        row_cells[1].text = str(st)
        row_cells[2].text = str(so)


    document.add_page_break()
        
    document.add_heading("Contact\n" , 1)
    p = document.add_paragraph('')
    #print(" ")
    m = py_get_int("ncbodys()")
    if m != 0:
        p.add_run(f"Number of Contact Bodys {m} \n\n")        
            
        table = document.add_table(rows=1, cols=2, style='Table Grid')
        #table.style = 'TableGrid'
        #table.style = 'ColorfulShading'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Contact Body Name'
        hdr_cells[1].text = 'ID'
            
        for i in range(1,m+1):
            sn = py_get_string("cbody_name_index(%d)" % i)
            id1 = py_get_int("cbody_id(%s)" % sn)
              
            row_cells = table.add_row().cells
            row_cells[0].text = str(sn)
            row_cells[1].text = str(id1)
          
          
        print("")
        p = document.add_paragraph('')
        sn = py_get_string("ctable_name()")
        p.add_run("\n\n  Contact Table     : " + str( sn) )
        sn = "contact_table:the_mesh:refined_mesh:dist_tol"
        dt = py_get_data(sn)
        p.add_run("  Contact Dist Tol  : " +str( dt) )
          
          
          
    else:
        p.add_run(f" No contact bodies \n")

    m = py_get_int("ngeoms()")
    print("\n Geometric Properties ",m)
    for i in range(1,m+1):
      sn = py_get_string("geom_name_index(%d)" % i)
      st = py_get_string("geom_type(%s)" % sn)
      p  = py_get_float("geom_par(%s,thick)" % sn)
      str1 = "  Geometric Prop  %12s  Type %19s   Thick %g" % (sn,st,p)
      print(str1)

    m = py_get_int("niconds()")
    print("\n Initial Conditions ",m)
    for i in range(1,m+1):
      sn = py_get_string("icond_name_index(%d)" % i)
      st = py_get_string("icond_type(%s)" % sn)
      so = py_get_string("icond_opt(%s,dof_values)" % sn)
      str1 = "  Initial  Cond %14s  Type %12s   Values by: %s" % (sn,st,so)
      print(str1)
      
    m = py_get_int("nsets()")
    print("Found ",m," sets")
    for i in range(1,m+1):
      id = py_get_int("set_id(%d)" % i)
      sn = py_get_string("set_name(%d)" % id)
      stype = py_get_string("set_type(%d)" % id)
      n = py_get_int("nset_entries(%d)" % id)

      if stype not in ("icond","apply","lcase"):
       print("Set ",sn,"is a ",stype," set with ",n," entries")
       """
       for j in range(1,n+1):
        k = py_get_int("set_entry(%d,%d)" % (id, j))
        print("    entry ",j," is ",k, end=' ')
        if (stype == 'face'):
          l = py_get_int("set_edge(%d,%d)" % (id, j))
          print("  face number ",l)
        elif( stype == 'edge'):
          l = py_get_int("set_edge(%d,%d)" % (id, j))
          print("  edge number ",l)
        else:
          print(" ")
        """

    print("")
    



    print("")
    sn = py_get_string("geom_name()")
    print("  Current geometry  : ", sn)
    thick = py_get_data("geometry:thick")
    print("  Thickness         : ", thick)

    """
    print("")
    m = py_get_int("nelements()")
    max_eid = py_get_int("max_element_id()")
    print(" Elements ", m, "   Maximum id ", max_eid)
    for i in range(1,m+1):
        id = py_get_int("element_id(%d)" % i)
        print("")
        sn = "element_class(%d)" % id
        e_cl = py_get_int(sn)
        sn = "element_family(%d)" % id
        e_fam = py_get_int(sn)
        sn = "element_type(%d)" % id
        e_ty = py_get_int(sn)
        print("  Element ", id, " Class ",e_cl," Family ",e_fam,", Type ",e_ty)
        cbn = py_get_string("element_cbody(%d)" % id)
        gmn = py_get_string("element_geom(%d)" % id)
        orn = py_get_string("element_orient(%d)" % id)
        mtn = py_get_string("element_mater(%d)" % id)
        print("    Contact Body      : ", cbn)
        print("    Geometry Property : ", gmn)
        print("    Orientation       : ", orn)
        print("    Material Property : ", mtn)
    """
    
    
    
    
    document.add_heading("Load Conditions\n" , 1)
    p = document.add_paragraph('')
    print("")
    sn = py_get_string("lcase_name_index(1)")
    st = py_get_string("lcase_type(%s)" % sn)
    p.add_run(f" Loadcase             : ,{sn}, , Type ,{st}\n")
    #st = py_get_string("lcase_opt(%s,arclength_meth)" % sn)
    #print("   ArcLength Method   : ",st)
    
    document.add_heading("JOB \n" , 1)
    p = document.add_paragraph('')
    p.add_run("\n")
    sn = py_get_string("job_name_index(1)")
    st = py_get_string("job_class(%s)" % sn)
    p.add_run(f" Job                  : {sn}, , Type ,{st} \n")

    
    #st = py_get_string("job_opt(%s,follow)" % sn)
    #print("   Follower Force     : ",st)




    """
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    
    
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    
    
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    
    
    document.add_page_break()
    
    """
    
    filename = "report_" + str(modelname1) + ".docx"

    document.save(filename)

    try:
        #py_send("*system_command %s" % filename)
        subprocess.Popen(['start', 'winword', filename], shell=True)
        pass
    except:
        print("no opending file")

    py_prompt("Done")

if __name__ == '__main__':
    py_connect('',40007)
    main()
    py_disconnect()